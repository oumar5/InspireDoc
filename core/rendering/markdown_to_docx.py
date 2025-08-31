import os
import re
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    raise ImportError("Veuillez installer python-docx: pip install python-docx")

try:
    import markdown
    from markdown.extensions import tables, fenced_code, codehilite
except ImportError:
    raise ImportError("Veuillez installer markdown: pip install markdown")

logger = logging.getLogger(__name__)

class MarkdownToDOCXConverter:
    """
    Convertisseur Markdown vers DOCX avec préservation de la mise en forme.
    """
    
    def __init__(self):
        """
        Initialise le convertisseur.
        """
        self.document = None
        self._init_styles_config()
    
    def _init_styles_config(self):
        """
        Initialise la configuration des styles.
        """
        self.styles_config = {
            'heading_1': {
                'font_size': Pt(18),
                'bold': True,
                'color': '2C3E50'
            },
            'heading_2': {
                'font_size': Pt(16),
                'bold': True,
                'color': '34495E'
            },
            'heading_3': {
                'font_size': Pt(14),
                'bold': True,
                'color': '34495E'
            },
            'normal': {
                'font_size': Pt(11),
                'font_name': 'Arial'
            },
            'code': {
                'font_size': Pt(10),
                'font_name': 'Courier New',
                'background_color': 'F1F2F6'
            }
        }
    
    def convert(self, 
                markdown_content: str,
                output_path: str,
                metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Convertit du contenu Markdown en DOCX.
        
        Args:
            markdown_content: Contenu Markdown à convertir
            output_path: Chemin de sortie pour le DOCX
            metadata: Métadonnées du document
            
        Returns:
            Dictionnaire avec le résultat de la conversion
        """
        try:
            # Créer un nouveau document
            self.document = Document()
            
            # Configurer les styles
            self._setup_document_styles()
            
            # Ajouter les métadonnées si fournies
            if metadata:
                self._add_document_properties(metadata)
            
            # Parser et convertir le Markdown
            self._parse_and_convert_markdown(markdown_content)
            
            # Sauvegarder le document
            self.document.save(output_path)
            
            # Métadonnées de conversion
            conversion_metadata = {
                'converted_at': datetime.now().isoformat(),
                'output_path': output_path,
                'file_size': os.path.getsize(output_path) if os.path.exists(output_path) else 0,
                'markdown_length': len(markdown_content),
                'paragraphs_count': len(self.document.paragraphs),
                'tables_count': len(self.document.tables)
            }
            
            if metadata:
                conversion_metadata['document_metadata'] = metadata
            
            logger.info(f"DOCX généré avec succès: {output_path} ({conversion_metadata['file_size']} bytes)")
            
            return {
                'success': True,
                'output_path': output_path,
                'metadata': conversion_metadata
            }
            
        except Exception as e:
            logger.error(f"Erreur lors de la conversion DOCX: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'output_path': output_path
            }
    
    def _setup_document_styles(self):
        """
        Configure les styles du document.
        """
        styles = self.document.styles
        
        # Style pour le code inline
        try:
            code_style = styles.add_style('Code Inline', WD_STYLE_TYPE.CHARACTER)
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(10)
        except Exception:
            pass  # Le style existe peut-être déjà
    
    def _add_document_properties(self, metadata: Dict[str, Any]):
        """
        Ajoute les propriétés du document.
        
        Args:
            metadata: Métadonnées à ajouter
        """
        core_props = self.document.core_properties
        
        if 'title' in metadata:
            core_props.title = metadata['title']
        if 'author' in metadata:
            core_props.author = metadata['author']
        if 'subject' in metadata:
            core_props.subject = metadata['subject']
        
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
    
    def _parse_and_convert_markdown(self, markdown_content: str):
        """
        Parse et convertit le contenu Markdown.
        
        Args:
            markdown_content: Contenu Markdown à parser
        """
        # Diviser le contenu en lignes
        lines = markdown_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Titres
            if line.startswith('#'):
                self._add_heading(line)
            
            # Listes
            elif line.strip().startswith(('-', '*', '+')) or re.match(r'^\s*\d+\.', line):
                i = self._add_list(lines, i)
                continue
            
            # Blocs de code
            elif line.strip().startswith('```'):
                i = self._add_code_block(lines, i)
                continue
            
            # Citations
            elif line.strip().startswith('>'):
                i = self._add_blockquote(lines, i)
                continue
            
            # Tableaux
            elif '|' in line and line.strip():
                i = self._add_table(lines, i)
                continue
            
            # Paragraphe normal
            elif line.strip():
                self._add_paragraph(line)
            
            # Ligne vide - ajouter un saut
            else:
                if self.document.paragraphs and self.document.paragraphs[-1].text.strip():
                    self.document.add_paragraph()
            
            i += 1
    
    def _add_heading(self, line: str):
        """
        Ajoute un titre.
        
        Args:
            line: Ligne contenant le titre
        """
        # Compter les # pour déterminer le niveau
        level = 0
        for char in line:
            if char == '#':
                level += 1
            else:
                break
        
        # Extraire le texte du titre
        title_text = line[level:].strip()
        
        # Ajouter le titre (limiter à 3 niveaux)
        heading_level = min(level, 3)
        heading = self.document.add_heading(title_text, level=heading_level)
        
        # Appliquer le style personnalisé
        if heading_level == 1:
            self._apply_heading_style(heading, 'heading_1')
        elif heading_level == 2:
            self._apply_heading_style(heading, 'heading_2')
        else:
            self._apply_heading_style(heading, 'heading_3')
    
    def _apply_heading_style(self, heading, style_name: str):
        """
        Applique un style à un titre.
        
        Args:
            heading: Élément titre
            style_name: Nom du style à appliquer
        """
        style_config = self.styles_config.get(style_name, {})
        
        for run in heading.runs:
            if 'font_size' in style_config:
                run.font.size = style_config['font_size']
            if 'bold' in style_config:
                run.font.bold = style_config['bold']
    
    def _add_paragraph(self, line: str):
        """
        Ajoute un paragraphe avec formatage inline.
        
        Args:
            line: Ligne de texte
        """
        paragraph = self.document.add_paragraph()
        
        # Parser le formatage inline
        self._parse_inline_formatting(paragraph, line)
    
    def _parse_inline_formatting(self, paragraph, text: str):
        """
        Parse le formatage inline (gras, italique, code).
        
        Args:
            paragraph: Paragraphe où ajouter le texte
            text: Texte à parser
        """
        # Patterns pour le formatage
        patterns = [
            (r'\*\*(.+?)\*\*', 'bold'),      # Gras
            (r'\*(.+?)\*', 'italic'),        # Italique
            (r'`(.+?)`', 'code'),            # Code inline
            (r'_(.+?)_', 'italic'),          # Italique alternatif
        ]
        
        remaining_text = text
        
        while remaining_text:
            earliest_match = None
            earliest_pos = len(remaining_text)
            
            # Trouver le premier pattern
            for pattern, format_type in patterns:
                match = re.search(pattern, remaining_text)
                if match and match.start() < earliest_pos:
                    earliest_match = (match, format_type)
                    earliest_pos = match.start()
            
            if earliest_match:
                match, format_type = earliest_match
                
                # Ajouter le texte avant le match
                if match.start() > 0:
                    paragraph.add_run(remaining_text[:match.start()])
                
                # Ajouter le texte formaté
                formatted_run = paragraph.add_run(match.group(1))
                
                if format_type == 'bold':
                    formatted_run.bold = True
                elif format_type == 'italic':
                    formatted_run.italic = True
                elif format_type == 'code':
                    formatted_run.font.name = 'Courier New'
                    formatted_run.font.size = Pt(10)
                
                # Continuer avec le reste du texte
                remaining_text = remaining_text[match.end():]
            else:
                # Pas de formatage trouvé, ajouter le reste
                paragraph.add_run(remaining_text)
                break
    
    def _add_list(self, lines: List[str], start_index: int) -> int:
        """
        Ajoute une liste.
        
        Args:
            lines: Toutes les lignes
            start_index: Index de début de la liste
            
        Returns:
            Index de la prochaine ligne à traiter
        """
        i = start_index
        
        while i < len(lines):
            line = lines[i]
            
            # Vérifier si c'est un élément de liste
            if line.strip().startswith(('-', '*', '+')):
                # Liste à puces
                text = line.strip()[1:].strip()
                paragraph = self.document.add_paragraph(text, style='List Bullet')
            elif re.match(r'^\s*\d+\.', line):
                # Liste numérotée
                text = re.sub(r'^\s*\d+\.\s*', '', line)
                paragraph = self.document.add_paragraph(text, style='List Number')
            else:
                # Fin de la liste
                break
            
            i += 1
        
        return i - 1
    
    def _add_code_block(self, lines: List[str], start_index: int) -> int:
        """
        Ajoute un bloc de code.
        
        Args:
            lines: Toutes les lignes
            start_index: Index de début du bloc
            
        Returns:
            Index de la prochaine ligne à traiter
        """
        i = start_index + 1  # Ignorer la ligne ```
        code_lines = []
        
        while i < len(lines):
            line = lines[i]
            if line.strip().startswith('```'):
                break
            code_lines.append(line)
            i += 1
        
        # Ajouter le bloc de code
        code_text = '\n'.join(code_lines)
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        return i
    
    def _add_blockquote(self, lines: List[str], start_index: int) -> int:
        """
        Ajoute une citation.
        
        Args:
            lines: Toutes les lignes
            start_index: Index de début de la citation
            
        Returns:
            Index de la prochaine ligne à traiter
        """
        i = start_index
        quote_lines = []
        
        while i < len(lines):
            line = lines[i]
            if line.strip().startswith('>'):
                quote_text = line.strip()[1:].strip()
                quote_lines.append(quote_text)
            else:
                break
            i += 1
        
        # Ajouter la citation
        quote_text = ' '.join(quote_lines)
        paragraph = self.document.add_paragraph(quote_text, style='Quote')
        
        return i - 1
    
    def _add_table(self, lines: List[str], start_index: int) -> int:
        """
        Ajoute un tableau.
        
        Args:
            lines: Toutes les lignes
            start_index: Index de début du tableau
            
        Returns:
            Index de la prochaine ligne à traiter
        """
        i = start_index
        table_lines = []
        
        while i < len(lines):
            line = lines[i]
            if '|' in line and line.strip():
                table_lines.append(line)
            else:
                break
            i += 1
        
        if len(table_lines) < 2:
            return i - 1
        
        # Parser le tableau
        rows = []
        for line in table_lines:
            if '---' in line:  # Ligne de séparation
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Ignorer les | de début/fin
            if cells:
                rows.append(cells)
        
        if not rows:
            return i - 1
        
        # Créer le tableau
        table = self.document.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        # Remplir le tableau
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_data
        
        return i - 1
    
    def convert_file(self, 
                    markdown_file_path: str,
                    output_path: Optional[str] = None,
                    metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Convertit un fichier Markdown en DOCX.
        
        Args:
            markdown_file_path: Chemin vers le fichier Markdown
            output_path: Chemin de sortie (optionnel)
            metadata: Métadonnées du document
            
        Returns:
            Dictionnaire avec le résultat de la conversion
        """
        try:
            # Lire le fichier Markdown
            with open(markdown_file_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # Générer le chemin de sortie si non fourni
            if not output_path:
                input_path = Path(markdown_file_path)
                output_path = str(input_path.with_suffix('.docx'))
            
            # Métadonnées du fichier
            file_metadata = metadata or {}
            file_metadata.update({
                'source_file': markdown_file_path,
                'file_size': os.path.getsize(markdown_file_path)
            })
            
            return self.convert(markdown_content, output_path, file_metadata)
            
        except Exception as e:
            logger.error(f"Erreur lors de la conversion du fichier {markdown_file_path}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'source_file': markdown_file_path
            }