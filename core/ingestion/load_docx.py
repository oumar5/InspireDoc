import logging
from typing import Dict, Any, List
from pathlib import Path

try:
    from docx import Document
except ImportError:
    raise ImportError("Veuillez installer python-docx: pip install python-docx")

logger = logging.getLogger(__name__)

class DOCXLoader:
    """
    Classe pour charger et extraire le texte des fichiers DOCX.
    """
    
    def __init__(self, include_headers: bool = True, include_footers: bool = True):
        """
        Initialise le loader DOCX.
        
        Args:
            include_headers: Inclure les en-têtes dans l'extraction
            include_footers: Inclure les pieds de page dans l'extraction
        """
        self.include_headers = include_headers
        self.include_footers = include_footers
        
    def load(self, file_path: str) -> Dict[str, Any]:
        """
        Charge un fichier DOCX et extrait son contenu.
        
        Args:
            file_path: Chemin vers le fichier DOCX
            
        Returns:
            Dictionnaire contenant le texte extrait et les métadonnées
        """
        try:
            doc = Document(file_path)
            
            # Extraction du texte principal
            paragraphs_text = self._extract_paragraphs(doc)
            
            # Extraction des tableaux
            tables_text = self._extract_tables(doc)
            
            # Extraction des en-têtes et pieds de page
            headers_text = self._extract_headers(doc) if self.include_headers else []
            footers_text = self._extract_footers(doc) if self.include_footers else []
            
            # Combinaison de tout le texte
            all_text_parts = []
            
            if headers_text:
                all_text_parts.extend(headers_text)
                all_text_parts.append("\n" + "="*50 + " CONTENU PRINCIPAL " + "="*50 + "\n")
            
            all_text_parts.extend(paragraphs_text)
            
            if tables_text:
                all_text_parts.append("\n" + "="*50 + " TABLEAUX " + "="*50 + "\n")
                all_text_parts.extend(tables_text)
            
            if footers_text:
                all_text_parts.append("\n" + "="*50 + " PIEDS DE PAGE " + "="*50 + "\n")
                all_text_parts.extend(footers_text)
            
            full_text = "\n".join(all_text_parts)
            
            # Métadonnées
            file_stats = Path(file_path).stat()
            metadata = {
                "file_path": file_path,
                "file_size_bytes": file_stats.st_size,
                "total_characters": len(full_text),
                "total_paragraphs": len(paragraphs_text),
                "total_tables": len(tables_text),
                "has_headers": len(headers_text) > 0,
                "has_footers": len(footers_text) > 0,
                "loader": "docx"
            }
            
            # Extraction des propriétés du document si disponibles
            try:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else ""
                })
            except Exception as e:
                logger.warning(f"Impossible d'extraire les propriétés du document: {str(e)}")
            
            logger.info(f"DOCX chargé: {metadata['total_paragraphs']} paragraphes, {metadata['total_tables']} tableaux")
            
            return {
                "text": full_text,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Erreur lors du chargement du DOCX {file_path}: {str(e)}")
            return {
                "text": "",
                "metadata": {
                    "error": str(e),
                    "file_path": file_path,
                    "total_characters": 0,
                    "total_paragraphs": 0
                }
            }
    
    def _extract_paragraphs(self, doc: Document) -> List[str]:
        """
        Extrait le texte des paragraphes.
        
        Args:
            doc: Document DOCX
            
        Returns:
            Liste des textes des paragraphes
        """
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Ignorer les paragraphes vides
                paragraphs.append(text)
        return paragraphs
    
    def _extract_tables(self, doc: Document) -> List[str]:
        """
        Extrait le texte des tableaux.
        
        Args:
            doc: Document DOCX
            
        Returns:
            Liste des textes des tableaux formatés
        """
        tables_text = []
        
        for table_idx, table in enumerate(doc.tables, 1):
            table_content = []
            table_content.append(f"\n--- Tableau {table_idx} ---")
            
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_cells.append(cell_text)
                
                # Joindre les cellules avec des séparateurs
                row_text = " | ".join(row_cells)
                if row_text.strip():  # Ignorer les lignes vides
                    table_content.append(row_text)
            
            if len(table_content) > 1:  # Si le tableau a du contenu
                tables_text.append("\n".join(table_content))
        
        return tables_text
    
    def _extract_headers(self, doc: Document) -> List[str]:
        """
        Extrait le texte des en-têtes.
        
        Args:
            doc: Document DOCX
            
        Returns:
            Liste des textes des en-têtes
        """
        headers = []
        
        try:
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        headers.append(f"[EN-TÊTE] {text}")
        except Exception as e:
            logger.warning(f"Erreur lors de l'extraction des en-têtes: {str(e)}")
        
        return headers
    
    def _extract_footers(self, doc: Document) -> List[str]:
        """
        Extrait le texte des pieds de page.
        
        Args:
            doc: Document DOCX
            
        Returns:
            Liste des textes des pieds de page
        """
        footers = []
        
        try:
            for section in doc.sections:
                footer = section.footer
                for paragraph in footer.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        footers.append(f"[PIED DE PAGE] {text}")
        except Exception as e:
            logger.warning(f"Erreur lors de l'extraction des pieds de page: {str(e)}")
        
        return footers
    
    @staticmethod
    def is_valid_docx(file_path: str) -> bool:
        """
        Vérifie si un fichier est un DOCX valide.
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            True si le fichier est un DOCX valide
        """
        try:
            Document(file_path)
            return True
        except Exception:
            return False